"""
교대 근무 인수인계 보고서 - Streamlit 앱
7개 질문에 답하면 Claude(Anthropic API)가 내용을 정리하고,
특이사항에 대한 대응 방법을 추천해서 인수인계 보고서를 만들어줍니다.
"""

import io
import re
from datetime import datetime

import streamlit as st
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ----------------------------------------------------------------------------
# 기본 설정
# ----------------------------------------------------------------------------
st.set_page_config(
    page_title="인수인계 보고서 작성",
    page_icon="📋",
    layout="centered",
)

ACCENT = "#1e3a5f"
CTA = "#2c7a7b"

MODEL_NAME = "claude-sonnet-5"

QUESTIONS = [
    ("q1", "Q1", "오늘 담당한 작업은 무엇인가요?", True),
    ("q2", "Q2", "작업 진행 상황은 어떠했나요?", False),
    ("q3", "Q3", "특이사항이 있었나요?", False),
    ("q4", "Q4", "문제가 발생했다면 어떤 내용이었나요?", False),
    ("q5", "Q5", "문제에 대해 어떤 조치를 했나요?", False),
    ("q6", "Q6", "다음 근무자가 이어서 해야 할 작업은 무엇인가요?", False),
    ("q7", "Q7", "다음 근무자가 주의해야 할 사항은 무엇인가요?", False),
]

# ----------------------------------------------------------------------------
# 스타일 (원본 html 카드 디자인 느낌을 살림)
# ----------------------------------------------------------------------------
st.markdown(
    f"""
    <style>
    .block-container {{ max-width: 680px; padding-top: 2.5rem; }}
    .eyebrow {{
        font-family: ui-monospace, "SF Mono", Consolas, monospace;
        font-size: 12px; letter-spacing: 0.13em; text-transform: uppercase;
        color: #8996a3; margin-bottom: 4px;
    }}
    h1 {{ font-weight: 800 !important; letter-spacing: -0.01em; }}
    .stButton>button[kind="primary"] {{
        background-color: {CTA}; border-color: {CTA};
    }}
    .stButton>button[kind="primary"]:hover {{
        background-color: #256665; border-color: #256665;
    }}
    .summary-line {{
        font-size: 15px; line-height: 1.6; background: #bdd6fc22;
        border-radius: 10px; padding: 12px 14px; border: 1px solid #bdd6fc66;
    }}
    </style>
    """,
    unsafe_allow_html=True,
)

# ----------------------------------------------------------------------------
# 세션 상태
# ----------------------------------------------------------------------------
if "stage" not in st.session_state:
    st.session_state.stage = "intro"  # intro -> form -> report
if "report_text" not in st.session_state:
    st.session_state.report_text = ""
if "meta" not in st.session_state:
    st.session_state.meta = {}


def get_client() -> Anthropic:
    api_key = st.secrets.get("ANTHROPIC_API_KEY", None)
    if not api_key:
        st.error(
            "ANTHROPIC_API_KEY가 설정되지 않았습니다. "
            "Streamlit Cloud의 App settings → Secrets에 "
            '`ANTHROPIC_API_KEY = "sk-ant-..."` 를 추가해주세요.'
        )
        st.stop()
    return Anthropic(api_key=api_key)


# ----------------------------------------------------------------------------
# Claude 호출 - 설문 내용을 보고서로 정리
# ----------------------------------------------------------------------------
def build_user_prompt(meta: dict, answers: dict) -> str:
    lines = [
        "아래 설문 내용을 종합·정리해서 인수인계 보고서로 작성해줘. "
        "특이사항이나 문제가 있다면, 그에 대응할 수 있는 방법을 3가지 정도 추천해줘. "
        "특이사항이나 문제가 없다면 대응 방법 섹션은 생략해줘.",
        "",
        f"작성자: {meta['writer']} / 근무조: {meta['shift']} / 작성일시: {meta['date_str']} {meta['time_str']}",
        "",
    ]
    for key, qnum, question, _ in QUESTIONS:
        answer = answers.get(key, "").strip() or "(답변 없음)"
        lines.append(f"[{qnum}] {question}")
        lines.append(answer)
        lines.append("")
    return "\n".join(lines)


SYSTEM_PROMPT = """당신은 제조 현장의 교대 근무 인수인계 보고서 작성을 돕는 전문 어시스턴트입니다.
사용자가 제공한 설문 응답을 바탕으로, 한국어로 마크다운 형식의 인수인계 보고서를 작성하세요.

보고서는 다음 구조를 따르세요:
## 1. 금일 담당 업무
## 2. 작업 진행 현황
## 3. 특이사항 및 문제 내용 (특이사항이 없으면 이 섹션에 "특이사항 없음"이라고만 적으세요)
## 4. 조치 내용 (조치한 내용이 없으면 이 섹션은 생략하세요)
## 5. 차기 근무자 인계사항
### 이어서 진행할 작업
### 주의사항
## 6. 특이사항 대응 방안 제안 (특이사항/문제가 있는 경우에만 작성, 3가지 방안을 ①②③ 형식으로, 각 방안은 제목과 1~2문장의 설명으로 구성)

규칙:
- 설문에 없는 내용을 지어내지 말고, 응답 내용을 자연스러운 문장으로 다듬어서 정리하세요.
- 불필요한 서론이나 마무리 인사 없이 보고서 본문만 출력하세요.
- 목록이 필요한 곳에는 "- " 로 시작하는 불릿을 사용하세요.
- 특이사항/문제가 없다면 4번, 6번 섹션은 출력하지 마세요.
"""


def generate_report(meta: dict, answers: dict) -> str:
    client = get_client()
    user_prompt = build_user_prompt(meta, answers)
    with st.spinner("Claude가 보고서를 정리하고 있어요..."):
        response = client.messages.create(
            model=MODEL_NAME,
            max_tokens=2000,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_prompt}],
        )
    text_blocks = [block.text for block in response.content if block.type == "text"]
    return "\n".join(text_blocks).strip()


# ----------------------------------------------------------------------------
# 마크다운 -> docx 변환 (간단한 헤딩 / 불릿 / 본문 처리)
# ----------------------------------------------------------------------------
def markdown_to_docx(markdown_text: str, meta: dict) -> bytes:
    doc = Document()

    title = doc.add_heading("생산 라인 인수인계 보고서", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.add_run(
        f"작성자: {meta['writer']}    근무조: {meta['shift']}    "
        f"작성일시: {meta['date_str']} {meta['time_str']}"
    ).italic = True

    doc.add_paragraph()

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("## "):
            heading = doc.add_heading(line[3:].strip(), level=1)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif re.match(r"^[-•]\s+", line):
            text = re.sub(r"^[-•]\s+", "", line)
            doc.add_paragraph(text, style="List Bullet")
        elif re.match(r"^①|^②|^③|^④|^⑤", line):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        else:
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            p = doc.add_paragraph(clean)
            for run in p.runs:
                run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ----------------------------------------------------------------------------
# 화면 1: 인트로
# ----------------------------------------------------------------------------
def render_intro():
    st.markdown('<p class="eyebrow">SHIFT HANDOVER · 교대 근무</p>', unsafe_allow_html=True)
    st.title("인수인계 보고서 작성")
    st.write(
        "7개 질문에 답하기만 하면, AI가 내용을 자연스럽게 정리하고 "
        "대응 방법까지 추천해서 보고서 파일로 만들어드려요."
    )
    st.markdown(
        """
- 📝 7개 질문에만 답하면 끝
- 🤖 AI가 내용을 정리하고 대응 방법까지 추천
- 📁 완성된 보고서를 파일(.docx)로 자동 생성
        """
    )
    if st.button("🚀 설문조사 시작하기", type="primary", use_container_width=True):
        st.session_state.stage = "form"
        st.rerun()


# ----------------------------------------------------------------------------
# 화면 2: 설문 폼
# ----------------------------------------------------------------------------
def render_form():
    st.markdown('<p class="eyebrow">SHIFT HANDOVER · 교대 근무</p>', unsafe_allow_html=True)
    st.title("인수인계 보고서 작성")
    st.write("아래 설문에 답하고 제출하면, Claude가 내용을 정리해서 보고서를 만들어드려요.")

    with st.form("handover_form"):
        col1, col2 = st.columns(2)
        with col1:
            writer = st.text_input("작성자", placeholder="이름")
        with col2:
            shift = st.selectbox("근무조", ["오전", "오후", "야간"])

        st.divider()

        answers = {}
        for key, qnum, question, required in QUESTIONS:
            label = f"**{qnum}**  {question}" + (" *(필수)*" if required else "")
            answers[key] = st.text_area(label, key=f"input_{key}", height=80)

        submitted = st.form_submit_button("📋 설문 제출하고 보고서 생성하기", type="primary", use_container_width=True)

    if submitted:
        if not answers["q1"].strip():
            st.error("⚠ Q1(오늘 담당한 작업)은 최소한 입력해주세요.")
            return

        now = datetime.now()
        meta = {
            "writer": writer.strip() or "미기재",
            "shift": shift,
            "date_str": now.strftime("%Y-%m-%d"),
            "time_str": now.strftime("%H:%M"),
            "file_date": now.strftime("%y%m%d"),
        }

        report_text = generate_report(meta, answers)

        st.session_state.meta = meta
        st.session_state.report_text = report_text
        st.session_state.stage = "report"
        st.rerun()


# ----------------------------------------------------------------------------
# 화면 3: 결과 리포트
# ----------------------------------------------------------------------------
def render_report():
    meta = st.session_state.meta
    st.markdown('<p class="eyebrow">SHIFT HANDOVER · 교대 근무</p>', unsafe_allow_html=True)
    st.title("📋 보고서가 완성되었어요")
    st.markdown(
        f'<div class="summary-line">작성일시: {meta["date_str"]} {meta["time_str"]} '
        f'· 작성자: {meta["writer"]} · 근무조: {meta["shift"]}</div>',
        unsafe_allow_html=True,
    )
    st.write("")
    st.markdown(st.session_state.report_text)

    st.divider()

    docx_bytes = markdown_to_docx(st.session_state.report_text, meta)
    filename = f"{meta['file_date']}_인수인계.docx"

    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            "📄 Word 파일(.docx) 다운로드",
            data=docx_bytes,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    with col2:
        if st.button("✏ 다시 작성하기", use_container_width=True):
            st.session_state.stage = "form"
            st.rerun()


# ----------------------------------------------------------------------------
# 라우팅
# ----------------------------------------------------------------------------
if st.session_state.stage == "intro":
    render_intro()
elif st.session_state.stage == "form":
    render_form()
else:
    render_report()
